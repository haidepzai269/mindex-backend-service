from __future__ import annotations

import os
import re
import tempfile
from dataclasses import asdict, dataclass
from urllib.parse import urlparse

import httpx

from app.handlers.base import BaseHandler, TaskError


def _validate_url(url: str) -> None:
    parsed = urlparse(url)
    if parsed.scheme not in ("http", "https") or not parsed.netloc:
        raise TaskError("URL must be a valid http/https URL")


@dataclass
class Block:
    type: str
    content: str
    page: int = 0
    level: int = 0


def normalize_text(text: str) -> str:
    return re.sub(r"\s+", " ", text or "").strip()


def is_bold_font(fontname: str) -> bool:
    lowered = (fontname or "").lower()
    return "bold" in lowered or "black" in lowered or "heavy" in lowered


def is_list_item(text: str) -> bool:
    return re.match(r"^(•|◦|-|\*|\d+\.|[A-Za-z]\.)\s+", text) is not None


def classify_line(
    text: str,
    page_num: int,
    line_size: float,
    avg_size: float,
    bold: bool,
    bottom: float | None = None,
    page_height: float | None = None,
) -> Block | None:
    text = normalize_text(text)
    if not text:
        return None

    if (
        bottom is not None
        and page_height is not None
        and len(text) < 10
        and text.isdigit()
        and abs(bottom - page_height) < 50
    ):
        return None

    if avg_size > 0 and (line_size > avg_size * 1.4 or (line_size > avg_size * 1.25 and bold)):
        return Block("heading1", text, page=page_num, level=1)
    if avg_size > 0 and line_size > avg_size * 1.15 and bold:
        return Block("heading2", text, page=page_num, level=2)
    if bold and (avg_size == 0 or line_size >= avg_size):
        if len(text) < 150:
            return Block("heading3", text, page=page_num, level=3)
        return Block("paragraph", text, page=page_num)
    if is_list_item(text):
        return Block("list_item", text, page=page_num)
    return Block("paragraph", text, page=page_num)


def group_words_into_lines(words: list, y_tolerance: float = 3.0) -> list:
    if not words:
        return []
    lines = []
    current_line = [words[0]]
    for word in words[1:]:
        if abs(word["top"] - current_line[-1]["top"]) <= y_tolerance or abs(word["bottom"] - current_line[-1]["bottom"]) <= y_tolerance:
            current_line.append(word)
        else:
            lines.append(sorted(current_line, key=lambda w: w["x0"]))
            current_line = [word]
    if current_line:
        lines.append(sorted(current_line, key=lambda w: w["x0"]))
    return lines


def extract_pdf_with_pdfplumber(path: str) -> list[Block]:
    import pdfplumber

    blocks = []
    with pdfplumber.open(path) as pdf:
        for page_num, page in enumerate(pdf.pages, 1):
            words = page.extract_words(extra_attrs=["size", "fontname"])
            if not words:
                continue

            sizes = [w["size"] for w in words if w.get("size")]
            if not sizes:
                continue
            avg_size = sum(sizes) / len(sizes)

            lines = group_words_into_lines(words)
            for line in lines:
                text = " ".join(w["text"] for w in line)
                line_size = float(line[0].get("size", avg_size) or avg_size or 0)
                bottom = float(line[0].get("bottom", 0) or 0)
                fontname = line[0].get("fontname", "")
                block = classify_line(
                    text, page_num, line_size, float(avg_size or 0),
                    is_bold_font(fontname), bottom, float(page.height),
                )
                if block is not None:
                    blocks.append(block)

            tables = page.extract_tables()
            for table in tables:
                rows = []
                for row in table:
                    merged = " | ".join(str(cell).strip() if cell else "" for cell in row)
                    rows.append(merged)
                if rows:
                    blocks.append(Block("table", "\n".join(rows), page=page_num))
    return blocks


def extract_pdf_with_pymupdf(path: str) -> list[Block]:
    import fitz

    blocks = []
    with fitz.open(path) as pdf:
        for page_num, page in enumerate(pdf, 1):
            page_dict = page.get_text("dict", sort=True)
            lines = []

            for block in page_dict.get("blocks", []):
                if block.get("type") != 0:
                    continue
                for line in block.get("lines", []):
                    spans = [span for span in line.get("spans", []) if normalize_text(span.get("text", ""))]
                    if not spans:
                        continue
                    text = "".join(span.get("text", "") for span in spans)
                    size = max(float(span.get("size", 0) or 0) for span in spans)
                    bold = any(is_bold_font(span.get("font", "")) for span in spans)
                    bbox = line.get("bbox") or block.get("bbox") or (0, 0, 0, 0)
                    bottom = float(bbox[3]) if len(bbox) >= 4 else 0
                    lines.append({"text": text, "size": size, "bold": bold, "bottom": bottom})

            if not lines:
                continue

            sizes = [line["size"] for line in lines if line["size"]]
            avg_size = sum(sizes) / len(sizes) if sizes else 0

            for line in lines:
                block = classify_line(
                    line["text"], page_num, line["size"], avg_size,
                    line["bold"], line["bottom"], float(page.rect.height),
                )
                if block is not None:
                    blocks.append(block)
    return blocks


def extract_pdf(path: str) -> list[Block]:
    primary_error = None
    try:
        blocks = extract_pdf_with_pdfplumber(path)
        if blocks:
            return blocks
    except Exception as exc:
        primary_error = exc

    try:
        fallback_blocks = extract_pdf_with_pymupdf(path)
    except Exception as fallback_exc:
        if primary_error is not None:
            raise RuntimeError(f"pdfplumber failed: {primary_error}; pymupdf failed: {fallback_exc}") from fallback_exc
        raise

    if fallback_blocks:
        return fallback_blocks
    if primary_error is not None:
        raise RuntimeError(f"pdfplumber failed: {primary_error}; pymupdf returned no text")
    return []


def is_office_by_magic(path):
    with open(path, "rb") as f:
        return f.read(4).startswith(b"PK\x03\x04")


def is_pdf_by_magic(path):
    with open(path, "rb") as f:
        return f.read(4).startswith(b"%PDF")


def markdown_to_blocks(markdown: str) -> list[Block]:
    blocks = []
    lines = markdown.split("\n")
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()
        if not stripped:
            i += 1
            continue

        if stripped.startswith("### "):
            blocks.append(Block("heading3", stripped[4:].strip(), level=3))
        elif stripped.startswith("## "):
            blocks.append(Block("heading2", stripped[3:].strip(), level=2))
        elif stripped.startswith("# "):
            blocks.append(Block("heading1", stripped[2:].strip(), level=1))
        elif stripped.startswith("|") and "|" in stripped:
            table_lines = []
            while i < len(lines):
                row = lines[i].strip()
                if not (row.startswith("|") and "|" in row):
                    break
                if not re.match(r"^\|[\s\-:]+\|", row):
                    table_lines.append(row)
                i += 1
            if table_lines:
                blocks.append(Block("table", "\n".join(table_lines)))
            continue
        elif stripped.startswith("```"):
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith("```"):
                code_lines.append(lines[i])
                i += 1
            if code_lines:
                blocks.append(Block("code", "\n".join(code_lines)))
        elif is_list_item(stripped):
            blocks.append(Block("list_item", stripped))
        else:
            blocks.append(Block("paragraph", stripped))

        i += 1
    return blocks


def extract_office_with_markitdown(path: str) -> list[Block]:
    from markitdown import MarkItDown

    result = MarkItDown().convert(path)
    content = result.text_content or ""
    if len(content.strip()) < 20:
        return []
    return markdown_to_blocks(content)


def extract_docx(path: str) -> list[Block]:
    try:
        blocks = extract_office_with_markitdown(path)
        if blocks:
            return blocks
    except Exception:
        pass

    from docx import Document

    doc = Document(path)
    blocks = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        style_name = para.style.name if para.style else "normal"
        style = style_name.lower()

        if "heading 1" in style:
            blocks.append(Block("heading1", text, level=1))
        elif "heading 2" in style:
            blocks.append(Block("heading2", text, level=2))
        elif "heading 3" in style:
            blocks.append(Block("heading3", text, level=3))
        elif style in ("list paragraph", "list bullet", "list number"):
            blocks.append(Block("list_item", text))
        elif style == "caption":
            blocks.append(Block("caption", text))
        elif "title" in style:
            blocks.append(Block("heading1", text, level=1))
        elif "subtitle" in style:
            blocks.append(Block("heading2", text, level=2))
        else:
            if style == "normal" and len(text) < 100:
                is_bold_flag = all(run.bold for run in para.runs if run.text.strip())
                if is_bold_flag and len(para.runs) > 0:
                    blocks.append(Block("heading3", text, level=3))
                    continue

            if para.runs and para.runs[0].font and para.runs[0].font.name in ("Courier New", "Consolas", "Courier", "Monaco"):
                blocks.append(Block("code", text))
            else:
                blocks.append(Block("paragraph", text))

    for table in doc.tables:
        rows = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            rows.append(" | ".join(cells))
        table_text = "\n".join(rows)
        if table_text.strip():
            blocks.append(Block("table", table_text))

    return blocks


async def _download_file(url: str, dest: str) -> None:
    async with httpx.AsyncClient(timeout=90.0) as client:
        resp = await client.get(url)
        resp.raise_for_status()
        with open(dest, "wb") as f:
            f.write(resp.content)


def _extract_sync(tmp_path: str, file_extension: str) -> list[Block]:
    ext = file_extension.lower()
    if is_pdf_by_magic(tmp_path) or ext == ".pdf":
        return extract_pdf(tmp_path)
    if is_office_by_magic(tmp_path) or ext in (".docx", ".doc", ".xlsx", ".pptx"):
        return extract_docx(tmp_path) if ext in (".docx", ".doc") else extract_office_with_markitdown(tmp_path)
    raise TaskError(f"Unsupported format: {ext}")


class ExtractHandler(BaseHandler):
    async def process(self, payload: dict) -> dict:
        import asyncio

        file_url = payload.get("file_url")
        file_extension = payload.get("file_extension", "")
        if not file_url:
            raise TaskError("Missing file_url in payload")
        _validate_url(file_url)

        tmp_dir = tempfile.mkdtemp(prefix="mindex_extract_")
        tmp_path = os.path.join(tmp_dir, f"document{file_extension}")

        try:
            await _download_file(file_url, tmp_path)

            loop = asyncio.get_running_loop()
            blocks = await loop.run_in_executor(None, _extract_sync, tmp_path, file_extension)

            if not blocks:
                blocks = [Block("empty", "", level=0)]

            return {"blocks": [asdict(b) for b in blocks]}
        finally:
            try:
                os.remove(tmp_path)
                os.rmdir(tmp_dir)
            except OSError:
                pass
